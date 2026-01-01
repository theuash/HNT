import time
import json
from datetime import datetime, timedelta
import os
import sys
from collections import deque
import threading
from dataclasses import dataclass
from enum import Enum
import shutil

n=0

if not os.path.exists(f"Data"):
    os.mkdir("Data")


while os.path.exists(f"./Data/Tracking-data_{n+1}"):
    n += 1
os.mkdir(f"./Data/Tracking-data_{n+1}")  
os.mkdir(f"./Data/Tracking-data_{n+1}/screenshots")
tracking_data = f"./Data/Tracking-data_{n+1}/"
screenshots=f"./Data/Tracking-data_{n+1}/screenshots/"





# Check and install required packages
def check_and_install_packages():
    """Check for required packages and install if missing"""
    required_packages = {
        'pyautogui': 'pyautogui',
        'psutil': 'psutil',
        'pynput': 'pynput',
        'docx': 'python-docx',
        'PIL': 'pillow'
    }
    
    missing_packages = []
    
    for import_name, package_name in required_packages.items():
        try:
            if import_name == 'docx':
                __import__('docx')
            else:
                __import__(import_name)
        except ImportError:
            missing_packages.append(package_name)
    
    if missing_packages:
        print("Missing packages detected.")
        print("Installing required packages...")
        import subprocess
        for package in missing_packages:
            print(f"Installing {package}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])
        print("\nPlease restart the script after installation completes.")
        return False
    return True

# Check packages before proceeding
if not check_and_install_packages():
    sys.exit(0)

# Now import the packages
import pyautogui
import psutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import ImageGrab
from pynput import keyboard, mouse
from pynput.mouse import Listener as MouseListener
from pynput.keyboard import Listener as KeyboardListener
import numpy as np

# Windows-specific imports
try:
    import win32gui
    import win32process
    WINDOWS_AVAILABLE = True
except ImportError:
    WINDOWS_AVAILABLE = False
    print("Note: Windows-specific features disabled. Basic tracking available.")

class ActivityType(Enum):
    MOUSE_MOVE = "Mouse Movement"
    MOUSE_CLICK = "Mouse Click"
    KEYBOARD_TYPING = "Keyboard Typing"
    WINDOW_SWITCH = "Window Switch"
    IDLE = "Idle"
    ACTIVE_WORK = "Active Work"

@dataclass
class MouseEvent:
    timestamp: datetime
    x: int
    y: int
    event_type: str  # 'move', 'click', 'scroll'
    button: str = None

@dataclass
class KeyboardEvent:
    timestamp: datetime
    key: str
    event_type: str  # 'press', 'release'

class AdvancedActivityTracker:

    def __init__(self, output_file="advanced_activity_report.docx", capture_screenshots=True):
        """
        Initialize the advanced activity tracker
        
        Args:
            output_file: Name of the Word document to create
            capture_screenshots: Whether to take periodic screenshots
        """
        self.output_file = output_file
        self.capture_screenshots = capture_screenshots
        self.activities = []
        self.start_time = None
        self.last_window = None
        self.screenshot_count = 0
        self.screenshots_dir = screenshots
        
        # Mouse tracking
        self.last_mouse_position = None
        self.mouse_movement_threshold = 5  # pixels
        self.mouse_events = deque(maxlen=1000)
        self.keyboard_events = deque(maxlen=1000)
        self.last_activity_time = None
        self.idle_start_time = None
        self.idle_periods = []
        
        # Activity state
        self.is_active = True
        self.active_start_time = None
        self.active_periods = []
        self.inactive_periods = []
        
        # Create directories if needed
        if not os.path.exists(self.screenshots_dir):
            os.makedirs(self.screenshots_dir)
        
        # Start listeners in separate threads
        self.mouse_listener = None
        self.keyboard_listener = None
        self.start_listeners()
    
    def start_listeners(self):
        """Start mouse and keyboard listeners"""
        # Mouse listener
        self.mouse_listener = MouseListener(
            on_move=self.on_mouse_move,
            on_click=self.on_mouse_click,
            on_scroll=self.on_mouse_scroll
        )
        self.mouse_listener.start()
        
        # Keyboard listener
        self.keyboard_listener = KeyboardListener(
            on_press=self.on_key_press,
            on_release=self.on_key_release
        )
        self.keyboard_listener.start()
        
        print("Mouse and keyboard listeners started...")
    
    def stop_listeners(self):
        """Stop all listeners"""
        if self.mouse_listener:
            self.mouse_listener.stop()
        if self.keyboard_listener:
            self.keyboard_listener.stop()
    
    def on_mouse_move(self, x, y):
        """Handle mouse movement events"""
        event = MouseEvent(
            timestamp=datetime.now(),
            x=x,
            y=y,
            event_type='move'
        )
        self.mouse_events.append(event)
        self.update_activity_state()
    
    def on_mouse_click(self, x, y, button, pressed):
        """Handle mouse click events"""
        if pressed:  # Only track when button is pressed, not released
            event = MouseEvent(
                timestamp=datetime.now(),
                x=x,
                y=y,
                event_type='click',
                button=str(button)
            )
            self.mouse_events.append(event)
            self.update_activity_state()
    
    def on_mouse_scroll(self, x, y, dx, dy):
        """Handle mouse scroll events"""
        event = MouseEvent(
            timestamp=datetime.now(),
            x=x,
            y=y,
            event_type='scroll'
        )
        self.mouse_events.append(event)
        self.update_activity_state()
    
    def on_key_press(self, key):
        """Handle keyboard key press events"""
        try:
            key_str = key.char
        except AttributeError:
            key_str = str(key)
        
        event = KeyboardEvent(
            timestamp=datetime.now(),
            key=key_str,
            event_type='press'
        )
        self.keyboard_events.append(event)
        self.update_activity_state()
    
    def on_key_release(self, key):
        """Handle keyboard key release events"""
        pass  # We're only tracking presses for activity detection
    
    def update_activity_state(self):
        """Update user activity state based on recent events"""
        current_time = datetime.now()
        self.last_activity_time = current_time
        
        # Check if user was idle and now became active
        if not self.is_active:
            self.is_active = True
            if self.idle_start_time:
                idle_duration = (current_time - self.idle_start_time).total_seconds()
                self.idle_periods.append({
                    'start': self.idle_start_time,
                    'end': current_time,
                    'duration': idle_duration
                })
                self.idle_start_time = None
            
            self.active_start_time = current_time
        
        # Reset idle tracking
        self.idle_start_time = None
    
    def check_idle_state(self):
        """Check if user has been idle for too long"""
        current_time = datetime.now()
        idle_threshold = 30  # 30 seconds of no activity = idle
        
        if self.last_activity_time:
            idle_time = (current_time - self.last_activity_time).total_seconds()
            
            if idle_time >= idle_threshold and self.is_active:
                self.is_active = False
                if self.active_start_time:
                    active_duration = (current_time - self.active_start_time).total_seconds()
                    self.active_periods.append({
                        'start': self.active_start_time,
                        'end': current_time,
                        'duration': active_duration
                    })
                
                self.idle_start_time = self.last_activity_time
    
    def get_active_window_info(self):
        """Get information about the currently active window"""
        try:
            if WINDOWS_AVAILABLE:
                window = win32gui.GetForegroundWindow()
                pid = win32process.GetWindowThreadProcessId(window)[1]
                
                try:
                    process = psutil.Process(pid)
                    exe_name = process.name()
                except:
                    exe_name = "Unknown"
                
                window_title = win32gui.GetWindowText(window)
                
                return {
                    'window_title': window_title,
                    'process_name': exe_name,
                    'pid': pid,
                    'time': datetime.now().strftime("%H:%M:%S")
                }
        except Exception:
            pass
        
        return {
            'window_title': "Active Window",
            'process_name': 'Unknown',
            'pid': 0,
            'time': datetime.now().strftime("%H:%M:%S")
        }
    
    def get_mouse_activity_summary(self, seconds=5):
        """Get mouse activity summary for the last N seconds"""
        current_time = datetime.now()
        threshold_time = current_time - timedelta(seconds=seconds)
        
        # Filter recent mouse events
        recent_moves = [e for e in self.mouse_events 
                       if e.timestamp > threshold_time and e.event_type == 'move']
        recent_clicks = [e for e in self.mouse_events 
                        if e.timestamp > threshold_time and e.event_type == 'click']
        
        # Calculate movement distance
        movement_distance = 0
        if len(recent_moves) > 1:
            for i in range(1, len(recent_moves)):
                x1, y1 = recent_moves[i-1].x, recent_moves[i-1].y
                x2, y2 = recent_moves[i].x, recent_moves[i].y
                movement_distance += ((x2-x1)**2 + (y2-y1)**2)**0.5
        
        return {
            'move_count': len(recent_moves),
            'click_count': len(recent_clicks),
            'movement_distance': movement_distance,
            'is_mouse_active': len(recent_moves) > 10 or len(recent_clicks) > 0
        }
    
    def get_keyboard_activity_summary(self, seconds=5):
        """Get keyboard activity summary for the last N seconds"""
        current_time = datetime.now()
        threshold_time = current_time - timedelta(seconds=seconds)
        
        # Filter recent keyboard events
        recent_keys = [e for e in self.keyboard_events 
                      if e.timestamp > threshold_time]
        
        # Calculate typing speed (keys per minute)
        keys_per_minute = 0
        if len(recent_keys) > 1:
            time_span = (recent_keys[-1].timestamp - recent_keys[0].timestamp).total_seconds()
            if time_span > 0:
                keys_per_minute = (len(recent_keys) / time_span) * 60
        
        return {
            'key_count': len(recent_keys),
            'keys_per_minute': keys_per_minute,
            'is_keyboard_active': len(recent_keys) > 5
        }
    
    def get_user_activity_status(self):
        """Determine user's current activity status"""
        mouse_summary = self.get_mouse_activity_summary(seconds=10)
        keyboard_summary = self.get_keyboard_activity_summary(seconds=10)
        
        is_mouse_active = mouse_summary['is_mouse_active']
        is_keyboard_active = keyboard_summary['is_keyboard_active']
        
        # Determine activity type
        if not self.is_active:
            return ActivityType.IDLE
        elif is_keyboard_active and keyboard_summary['keys_per_minute'] > 20:
            return ActivityType.KEYBOARD_TYPING
        elif is_mouse_active and mouse_summary['movement_distance'] > 100:
            return ActivityType.MOUSE_MOVE
        elif is_mouse_active:
            return ActivityType.MOUSE_CLICK
        else:
            return ActivityType.ACTIVE_WORK if self.is_active else ActivityType.IDLE
    
    def get_focus_level(self):
        """Calculate user's focus level based on activity"""
        mouse_summary = self.get_mouse_activity_summary(seconds=30)
        keyboard_summary = self.get_keyboard_activity_summary(seconds=30)
        
        # Simple focus calculation (0-100)
        focus = 0
        
        if keyboard_summary['keys_per_minute'] > 30:
            focus += 40  # Fast typing indicates high focus
        elif keyboard_summary['keys_per_minute'] > 10:
            focus += 20
        
        if mouse_summary['movement_distance'] > 200:
            focus += 30  # Significant mouse movement
        elif mouse_summary['movement_distance'] > 50:
            focus += 15
        
        # Add points for being active
        if self.is_active:
            focus += 30
        
        return min(100, focus)
    
    def capture_screenshot(self):
        """Take a screenshot and save it"""
        if not self.capture_screenshots:
            return None
            
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{self.screenshots_dir}/screenshot_{timestamp}.png"
        
        try:
            screenshot = ImageGrab.grab()
            screenshot.save(filename, 'PNG')
            self.screenshot_count += 1
            return filename
        except Exception as e:
            print(f"Failed to capture screenshot: {e}")
            return None
    
    def get_system_info(self):
        """Get current system information"""
        try:
            mouse_x, mouse_y = pyautogui.position()
        except:
            mouse_x, mouse_y = 0, 0
        
        try:
            cpu_percent = psutil.cpu_percent()
            memory_percent = psutil.virtual_memory().percent
        except:
            cpu_percent = 0
            memory_percent = 0
        
        return {
            'mouse_position': f"({mouse_x}, {mouse_y})",
            'time': datetime.now().strftime("%H:%M:%S %Y-%m-%d"),
            'cpu_percent': cpu_percent,
            'memory_percent': memory_percent
        }
    
    def track(self, duration_seconds=300, interval_seconds=5):
        """
        Track activity for specified duration
        
        Args:
            duration_seconds: How long to track (in seconds)
            interval_seconds: How often to check (in seconds)
        """
        print(f"Starting advanced activity tracking for {duration_seconds} seconds...")
        print("Tracking: Mouse movements, clicks, keyboard typing, window focus")
        print("Press Ctrl+C to stop early.")
        print("-" * 50)
        
        self.start_time = time.time()
        self.last_activity_time = datetime.now()
        self.active_start_time = datetime.now()
        end_time = self.start_time + duration_seconds
        
        try:
            while time.time() < end_time:
                # Check idle state
                self.check_idle_state()
                
                # Get current window info
                window_info = self.get_active_window_info()
                
                # Get system info
                system_info = self.get_system_info()
                
                # Get activity summaries
                mouse_summary = self.get_mouse_activity_summary()
                keyboard_summary = self.get_keyboard_activity_summary()
                activity_status = self.get_user_activity_status()
                focus_level = self.get_focus_level()
                
                # Capture screenshot if enabled
                screenshot_path = None
                if self.capture_screenshots:
                    screenshot_path = self.capture_screenshot()
                
                # Check if window changed
                current_window = window_info['window_title']
                window_changed = (current_window != self.last_window) if self.last_window else True
                
                if window_changed or interval_seconds >= 5:
                    activity = {
                        'timestamp': datetime.now().isoformat(),
                        'window_title': window_info['window_title'],
                        'process': window_info['process_name'],
                        'mouse_position': system_info['mouse_position'],
                        'cpu_usage': system_info['cpu_percent'],
                        'memory_usage': system_info['memory_percent'],
                        'mouse_moves': mouse_summary['move_count'],
                        'mouse_clicks': mouse_summary['click_count'],
                        'mouse_distance': mouse_summary['movement_distance'],
                        'keyboard_keys': keyboard_summary['key_count'],
                        'typing_speed': keyboard_summary['keys_per_minute'],
                        'activity_status': activity_status.value,
                        'focus_level': focus_level,
                        'is_user_active': self.is_active,
                        'idle_duration': self.get_current_idle_duration(),
                        'screenshot': screenshot_path,
                        'type': 'window_change' if window_changed else 'periodic'
                    }
                    self.activities.append(activity)
                    self.last_window = current_window
                    
                    # Display status
                    status_symbol = "‚úì" if self.is_active else "‚úó"
                    print(f"[{system_info['time']}] {status_symbol} {activity_status.value} | "
                          f"Mouse: {mouse_summary['move_count']}m/{mouse_summary['click_count']}c | "
                          f"Keys: {keyboard_summary['key_count']} | "
                          f"Focus: {focus_level}%")
                
                time.sleep(interval_seconds)
                
        except KeyboardInterrupt:
            print("\nTracking stopped by user.")
        except Exception as e:
            print(f"Error during tracking: {e}")
            import traceback
            traceback.print_exc()
        
        # Record final active period if still active
        if self.is_active and self.active_start_time:
            active_duration = (datetime.now() - self.active_start_time).total_seconds()
            self.active_periods.append({
                'start': self.active_start_time,
                'end': datetime.now(),
                'duration': active_duration
            })
        
        print(f"\nTracking completed. Collected {len(self.activities)} activity entries.")
        print(f"Active periods: {len(self.active_periods)}")
        print(f"Idle periods: {len(self.idle_periods)}")
    
    def get_current_idle_duration(self):
        """Get current idle duration in seconds"""
        if self.idle_start_time:
            return (datetime.now() - self.idle_start_time).total_seconds()
        return 0
    
    def create_word_document(self):
        """Create a comprehensive Word document with detailed activity analysis"""
        print(f"\nCreating comprehensive Word document: {self.output_file}")
        
        doc = Document()
        
        # Title with formatting
        title = doc.add_heading('Advanced Activity Tracking Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary information
        doc.add_heading('Executive Summary', level=1)
        summary = doc.add_paragraph()
        summary.add_run(f"Tracking Period: ").bold = True
        tracking_duration = (datetime.now() - datetime.fromtimestamp(self.start_time)).total_seconds() / 60
        summary.add_run(f"{tracking_duration:.1f} minutes\n")
        
        summary.add_run(f"Total Activity Entries: ").bold = True
        summary.add_run(f"{len(self.activities)}\n")
        
        # Calculate statistics
        if self.activities:
            total_active_time = sum(p['duration'] for p in self.active_periods)
            total_idle_time = sum(p['duration'] for p in self.idle_periods)
            total_time = total_active_time + total_idle_time
            
            productivity_score = self.calculate_productivity_score()
            
            summary.add_run(f"Total Active Time: ").bold = True
            summary.add_run(f"{total_active_time/60:.1f} minutes ({total_active_time/total_time*100:.1f}%)\n")
            
            summary.add_run(f"Total Idle Time: ").bold = True
            summary.add_run(f"{total_idle_time/60:.1f} minutes ({total_idle_time/total_time*100:.1f}%)\n")
            
            summary.add_run(f"Productivity Score: ").bold = True
            summary.add_run(f"{productivity_score}/100\n")
        
        summary.add_run(f"Report Generated: ").bold = True
        summary.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        
        # Productivity Analysis Section
        doc.add_heading('Productivity Analysis', level=1)
        
        if self.active_periods and self.idle_periods:
            # Create productivity table
            prod_table = doc.add_table(rows=1, cols=4)
            prod_table.style = 'LightShading'
            hdr_cells = prod_table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Value'
            hdr_cells[2].text = 'Duration'
            hdr_cells[3].text = 'Percentage'
            
            metrics = [
                ('Active Work Periods', len(self.active_periods), 
                 sum(p['duration'] for p in self.active_periods)),
                ('Idle Periods', len(self.idle_periods),
                 sum(p['duration'] for p in self.idle_periods)),
            ]
            
            total_seconds = metrics[0][2] + metrics[1][2]
            
            for name, count, duration in metrics:
                row_cells = prod_table.add_row().cells
                row_cells[0].text = name
                row_cells[1].text = str(count)
                row_cells[2].text = f"{duration/60:.1f} min"
                row_cells[3].text = f"{duration/total_seconds*100:.1f}%" if total_seconds > 0 else "0%"
        
        # Activity Type Breakdown
        doc.add_heading('Activity Type Breakdown', level=1)
        if self.activities:
            activity_counts = {}
            for activity in self.activities:
                status = activity.get('activity_status', 'Unknown')
                activity_counts[status] = activity_counts.get(status, 0) + 1
            
            activity_table = doc.add_table(rows=1, cols=3)
            activity_table.style = 'LightGrid'
            hdr_cells = activity_table.rows[0].cells
            hdr_cells[0].text = 'Activity Type'
            hdr_cells[1].text = 'Count'
            hdr_cells[2].text = 'Percentage'
            
            total = len(self.activities)
            for activity_type, count in sorted(activity_counts.items(), key=lambda x: x[1], reverse=True):
                row_cells = activity_table.add_row().cells
                row_cells[0].text = activity_type
                row_cells[1].text = str(count)
                row_cells[2].text = f"{(count/total*100):.1f}%"
        
        # Applications Usage with Focus Analysis
        doc.add_heading('Application Usage Analysis', level=1)
        if self.activities:
            app_stats = {}
            for activity in self.activities:
                app = activity['process']
                if app not in app_stats:
                    app_stats[app] = {
                        'count': 0,
                        'total_focus': 0,
                        'window_titles': set()
                    }
                app_stats[app]['count'] += 1
                app_stats[app]['total_focus'] += activity.get('focus_level', 0)
                app_stats[app]['window_titles'].add(activity['window_title'][:50])
            
            app_table = doc.add_table(rows=1, cols=5)
            app_table.style = 'LightShading'
            hdr_cells = app_table.rows[0].cells
            hdr_cells[0].text = 'Application'
            hdr_cells[1].text = 'Usage Count'
            hdr_cells[2].text = 'Avg Focus %'
            hdr_cells[3].text = 'Usage %'
            hdr_cells[4].text = 'Windows Used'
            
            total = len(self.activities)
            for app, stats in sorted(app_stats.items(), key=lambda x: x[1]['count'], reverse=True)[:10]:
                row_cells = app_table.add_row().cells
                row_cells[0].text = app
                row_cells[1].text = str(stats['count'])
                avg_focus = stats['total_focus'] / stats['count'] if stats['count'] > 0 else 0
                row_cells[2].text = f"{avg_focus:.1f}%"
                row_cells[3].text = f"{(stats['count']/total*100):.1f}%"
                row_cells[4].text = ', '.join(list(stats['window_titles'])[:2]) + ('...' if len(stats['window_titles']) > 2 else '')
        
        # Detailed Activity Log
        doc.add_heading('Detailed Activity Log', level=1)
        doc.add_paragraph('This section contains detailed tracking of all activities.')
        
        if self.activities:
            log_table = doc.add_table(rows=1, cols=8)
            log_table.style = 'LightGrid'
            hdr_cells = log_table.rows[0].cells
            hdr_cells[0].text = 'Time'
            hdr_cells[1].text = 'Application'
            hdr_cells[2].text = 'Activity'
            hdr_cells[3].text = 'Mouse'
            hdr_cells[4].text = 'Keyboard'
            hdr_cells[5].text = 'Focus %'
            hdr_cells[6].text = 'Status'
            hdr_cells[7].text = 'Window'
            
            # Show last 50 entries for readability
            display_activities = self.activities[-50:] if len(self.activities) > 50 else self.activities
            
            for activity in display_activities:
                timestamp = datetime.fromisoformat(activity['timestamp']).strftime("%H:%M:%S")
                row_cells = log_table.add_row().cells
                row_cells[0].text = timestamp
                row_cells[1].text = activity['process'][:15]
                row_cells[2].text = activity['activity_status'][:15]
                row_cells[3].text = f"{activity['mouse_moves']}m/{activity['mouse_clicks']}c"
                row_cells[4].text = f"{activity['keyboard_keys']}k"
                row_cells[5].text = f"{activity['focus_level']}%"
                row_cells[6].text = "Active" if activity['is_user_active'] else "Idle"
                row_cells[7].text = activity['window_title'][:25] + "..." if len(activity['window_title']) > 25 else activity['window_title']
            
            if len(self.activities) > 50:
                doc.add_paragraph(f'Note: Showing last 50 of {len(self.activities)} total entries.')
        
        # Recommendations Section
        doc.add_heading('Productivity Recommendations', level=1)
        recommendations = doc.add_paragraph()
        
        if self.activities:
            avg_focus = sum(a.get('focus_level', 0) for a in self.activities) / len(self.activities)
            
            recommendations.add_run("Based on your activity patterns:\n\n").bold = True
            
            if avg_focus < 40:
                recommendations.add_run("‚ö†Ô∏è Your average focus level is low. Consider:\n")
                recommendations.add_run("‚Ä¢ Taking regular breaks using Pomodoro technique\n")
                recommendations.add_run("‚Ä¢ Minimizing distractions from non-work applications\n")
                recommendations.add_run("‚Ä¢ Setting specific goals for each work session\n\n")
            elif avg_focus < 70:
                recommendations.add_run("‚úÖ Your focus level is moderate. To improve:\n")
                recommendations.add_run("‚Ä¢ Track your most productive times of day\n")
                recommendations.add_run("‚Ä¢ Use focus-enhancing apps or browser extensions\n")
                recommendations.add_run("‚Ä¢ Practice single-tasking instead of multitasking\n\n")
            else:
                recommendations.add_run("üéâ Excellent focus level! Keep up the good work!\n")
                recommendations.add_run("‚Ä¢ Consider sharing your productivity strategies\n")
                recommendations.add_run("‚Ä¢ Take regular breaks to maintain this level\n")
                recommendations.add_run("‚Ä¢ Review what makes your most focused sessions successful\n\n")
            
            # Check for frequent switching
            window_switches = sum(1 for a in self.activities if a.get('type') == 'window_change')
            if window_switches > len(self.activities) * 0.3:  # More than 30% of activities are window switches
                recommendations.add_run("‚ö†Ô∏è You switch windows frequently. This reduces productivity.\n")
                recommendations.add_run("‚Ä¢ Try to batch similar tasks together\n")
                recommendations.add_run("‚Ä¢ Use virtual desktops to organize workspaces\n")
                recommendations.add_run("‚Ä¢ Turn off notifications during deep work sessions\n")
        
        # Save the document
        doc.save(f"{tracking_data}/{self.output_file}")
        print(f"Word document created: ./Data/tracking_data/{self.output_file}")
        
        # Save raw data as JSON
        json_file = f"{tracking_data}/{self.output_file.replace('.docx', '_data.json')}"
        output_data = {
            'activities': self.activities,
            'active_periods': [{
                'start': p['start'].isoformat(),
                'end': p['end'].isoformat(),
                'duration': p['duration']
            } for p in self.active_periods],
            'idle_periods': [{
                'start': p['start'].isoformat(),
                'end': p['end'].isoformat(),
                'duration': p['duration']
            } for p in self.idle_periods],
            'summary': {
                'total_activities': len(self.activities),
                'total_active_time': sum(p['duration'] for p in self.active_periods),
                'total_idle_time': sum(p['duration'] for p in self.idle_periods),
                'productivity_score': self.calculate_productivity_score()
            }
        }
        
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, indent=2, ensure_ascii=False)
        print(f"Raw data saved as: {json_file}")
    
    def calculate_productivity_score(self):
        """Calculate overall productivity score (0-100)"""
        if not self.activities:
            return 0
        
        # Weighted factors
        avg_focus = sum(a.get('focus_level', 0) for a in self.activities) / len(self.activities)
        
        # Active vs idle ratio
        total_active = sum(p['duration'] for p in self.active_periods)
        total_idle = sum(p['duration'] for p in self.idle_periods)
        total_time = total_active + total_idle
        
        if total_time == 0:
            activity_ratio = 0
        else:
            activity_ratio = (total_active / total_time) * 100
        
        # Keyboard activity indicates productive work
        avg_typing_speed = sum(a.get('typing_speed', 0) for a in self.activities) / len(self.activities)
        typing_score = min(100, avg_typing_speed * 2)  # Scale typing speed
        
        # Calculate final score (weighted average)
        score = (avg_focus * 0.4) + (activity_ratio * 0.4) + (typing_score * 0.2)
        
        return round(score, 1)
    
    def run(self, duration_minutes=5, capture_screenshots=False):
        """
        Run the complete tracking and reporting process
        
        Args:
            duration_minutes: How many minutes to track
            capture_screenshots: Whether to capture screenshots
        """
        self.capture_screenshots = capture_screenshots
        
        print("=" * 70)
        print("ADVANCED ACTIVITY TRACKER")
        print("=" * 70)
        print("This program will track in detail:")
        print("‚úì Active/Idle status with timing")
        print("‚úì Mouse movements, clicks, and distance traveled")
        print("‚úì Keyboard typing speed and activity")
        print("‚úì Window focus and switches")
        print("‚úì Real-time focus level calculation")
        print("‚úì Productivity analysis and recommendations")
        if capture_screenshots:
            print("‚úì Periodic screenshots")
        print(f"‚úì Create comprehensive report after {duration_minutes} minutes")
        print("=" * 70)
        
        
        try:
            # Start tracking
            self.track(duration_seconds=duration_minutes * 60)
            
            # Create report
            self.create_word_document()
            
            # Stop listeners
            self.stop_listeners()
            

            print("=" * 70)
            print("REPORT COMPLETE")
            print(f"1. Word Document: {self.output_file}")
            print(f"2. Raw Data: {self.output_file.replace('.docx', '_data.json')}")
            if capture_screenshots:
                print(f"3. Screenshots: {self.screenshots_dir}/")
            
            # Print summary
            if self.active_periods and self.idle_periods:
                total_active = sum(p['duration'] for p in self.active_periods)
                total_idle = sum(p['duration'] for p in self.idle_periods)
                productivity = self.calculate_productivity_score()
                print(f"\nSUMMARY:")
                print(f"‚Ä¢ Active Time: {total_active/60:.1f} minutes")
                print(f"‚Ä¢ Idle Time: {total_idle/60:.1f} minutes")
                print(f"‚Ä¢ Productivity Score: {productivity}/100")
            
            print("=" * 70)
            
        except Exception as e:
            print(f"Error during execution: {e}")
            import traceback
            traceback.print_exc()
            self.stop_listeners()


if __name__ == "__main__":
    try:
        # Create tracker instance
        timestamp = datetime.now().strftime('%Y%m%d_%H%M')
        tracker = AdvancedActivityTracker(
            output_file=f"advanced_activity_report_{timestamp}.docx"
        )
        
        # Configure tracking parameters
        TRACKING_DURATION_MINUTES = 120  # Change this as needed
        CAPTURE_SCREENSHOTS = True    # Set to True to enable screenshots
        
        tracker.run(
            duration_minutes=TRACKING_DURATION_MINUTES,
            capture_screenshots=CAPTURE_SCREENSHOTS
        )
        
    except KeyboardInterrupt:
        print("\nProgram cancelled by user.")
    except Exception as e:
        print(f"An error occurred: {e}")
        print("\nTroubleshooting tips:")
        print("1. Run as administrator if on Windows")
        print("2. Make sure all dependencies are installed")
        print("3. Check if antivirus is blocking the mouse/keyboard listeners")

    




